from __future__ import annotations

import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer


GUIDANCE_DIR = Path(__file__).resolve().parent
MARKDOWN_FILE = Path(os.getenv("GUIDE_MANUAL_MD", str(GUIDANCE_DIR / "accounting_manual.md")))
DOCX_FILE = Path(
    os.getenv(
        "GUIDE_MANUAL_DOCX",
        str(MARKDOWN_FILE.with_suffix(".docx")),
    )
)
PDF_FILE = Path(
    os.getenv(
        "GUIDE_MANUAL_PDF",
        str(MARKDOWN_FILE.with_suffix(".pdf")),
    )
)

IMAGE_PATTERN = re.compile(r"!\[(.*?)\]\((.*?)\)")


def register_unicode_font():
    candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/tahoma.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for font_path in candidates:
        if font_path.exists():
            pdfmetrics.registerFont(TTFont("GuideUnicode", str(font_path)))
            return "GuideUnicode"
    return "Helvetica"


def parse_markdown(text: str):
    blocks: list[tuple[str, str]] = []
    paragraph_lines: list[str] = []

    def flush_paragraph():
        if paragraph_lines:
            blocks.append(("p", " ".join(line.strip() for line in paragraph_lines if line.strip())))
            paragraph_lines.clear()

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            flush_paragraph()
            continue
        if line.startswith("# "):
            flush_paragraph()
            blocks.append(("h1", line[2:].strip()))
            continue
        if line.startswith("## "):
            flush_paragraph()
            blocks.append(("h2", line[3:].strip()))
            continue
        if line.strip() == "---":
            flush_paragraph()
            blocks.append(("pagebreak", ""))
            continue
        if line.startswith("- "):
            flush_paragraph()
            blocks.append(("li", line[2:].strip()))
            continue
        image_match = IMAGE_PATTERN.match(line.strip())
        if image_match:
            flush_paragraph()
            _, path_text = image_match.groups()
            blocks.append(("img", path_text))
            continue
        paragraph_lines.append(line)

    flush_paragraph()
    return blocks


def build_docx(blocks):
    document = Document()
    document.add_heading(MARKDOWN_FILE.stem.replace("_", " ").title(), level=0)

    for block_type, content in blocks:
        if block_type == "h1":
            document.add_heading(content, level=1)
        elif block_type == "h2":
            document.add_heading(content, level=2)
        elif block_type == "li":
            document.add_paragraph(content, style="List Bullet")
        elif block_type == "img":
            image_path = (GUIDANCE_DIR / content).resolve()
            if image_path.exists():
                document.add_picture(str(image_path), width=Inches(6.2))
            else:
                document.add_paragraph(f"[Chưa có ảnh: {content}]")
        elif block_type == "pagebreak":
            document.add_page_break()
        else:
            document.add_paragraph(content)

    document.save(DOCX_FILE)


def build_pdf(blocks):
    styles = getSampleStyleSheet()
    font_name = register_unicode_font()
    styles["Title"].fontName = font_name
    styles["Heading2"].fontName = font_name
    styles["BodyText"].fontName = font_name
    styles["Italic"].fontName = font_name
    story = []
    doc = SimpleDocTemplate(str(PDF_FILE), pagesize=A4, rightMargin=1.5 * cm, leftMargin=1.5 * cm)

    for block_type, content in blocks:
        if block_type == "h1":
            story.append(Paragraph(f"<b>{content}</b>", styles["Title"]))
            story.append(Spacer(1, 0.3 * cm))
        elif block_type == "h2":
            story.append(Paragraph(f"<b>{content}</b>", styles["Heading2"]))
            story.append(Spacer(1, 0.15 * cm))
        elif block_type == "li":
            story.append(Paragraph(f"• {content}", styles["BodyText"]))
            story.append(Spacer(1, 0.08 * cm))
        elif block_type == "img":
            image_path = (GUIDANCE_DIR / content).resolve()
            if image_path.exists():
                story.append(Image(str(image_path), width=17.5 * cm, height=10 * cm, kind="proportional"))
            else:
                story.append(Paragraph(f"[Chưa có ảnh: {content}]", styles["Italic"]))
            story.append(Spacer(1, 0.25 * cm))
        elif block_type == "pagebreak":
            story.append(PageBreak())
        else:
            story.append(Paragraph(content, styles["BodyText"]))
            story.append(Spacer(1, 0.15 * cm))

    doc.build(story)


def main():
    markdown_text = MARKDOWN_FILE.read_text(encoding="utf-8")
    blocks = parse_markdown(markdown_text)
    build_docx(blocks)
    build_pdf(blocks)
    sys.stdout.buffer.write(f"Đã tạo: {DOCX_FILE}\n".encode("utf-8"))
    sys.stdout.buffer.write(f"Đã tạo: {PDF_FILE}\n".encode("utf-8"))


if __name__ == "__main__":
    main()
